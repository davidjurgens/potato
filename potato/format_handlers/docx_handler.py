"""
DOCX Format Handler

Extracts text and structure from Word documents using python-docx and mammoth.
Supports paragraph-level coordinate mapping for span annotations.

Usage:
    from potato.format_handlers.docx_handler import DocxHandler

    handler = DocxHandler()
    output = handler.extract("document.docx", {
        "preserve_styles": True,
        "include_headers": True,
    })
"""

from typing import Dict, List, Any, Optional
from pathlib import Path
import html
import logging
import uuid

from .base import BaseFormatHandler, FormatOutput
from .coordinate_mapping import CoordinateMapper, DocumentCoordinate

logger = logging.getLogger(__name__)

# Check if dependencies are available
try:
    import docx
    from docx.document import Document
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    docx = None

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False
    mammoth = None


class DocxHandler(BaseFormatHandler):
    """
    Handler for Word documents (.docx).

    Uses python-docx for text extraction with structure preservation
    and mammoth for HTML conversion.
    """

    format_name = "docx"
    supported_extensions = [".docx"]
    description = "Word document extraction with paragraph/section mapping"
    requires_dependencies = ["python-docx", "mammoth"]

    def get_default_options(self) -> Dict[str, Any]:
        """Get default extraction options."""
        return {
            "preserve_styles": True,
            "include_headers": True,
            "include_footers": False,
            "include_tables": True,
            "paragraph_separator": "\n\n",
            "use_mammoth_html": True,  # Use mammoth for rich HTML conversion
        }

    def extract(
        self,
        file_path: str,
        options: Optional[Dict[str, Any]] = None
    ) -> FormatOutput:
        """
        Extract text and structure from a Word document.

        Args:
            file_path: Path to the .docx file
            options: Extraction options:
                - preserve_styles: Keep heading levels and formatting
                - include_headers: Include document headers
                - include_footers: Include document footers
                - include_tables: Extract table content

        Returns:
            FormatOutput with extracted text, HTML, and coordinate mappings
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX extraction. "
                "Install with: pip install python-docx"
            )

        opts = self.merge_options(options)
        mapper = CoordinateMapper()

        # Open document
        doc = docx.Document(file_path)

        # Extract text with structure
        text_parts = []
        html_parts = []
        current_offset = 0

        metadata = {
            "format": "docx",
            "source_file": str(file_path),
            "paragraphs": [],
            "sections": [],
            "tables": [],
        }

        # Check if mammoth is available for rich HTML
        if opts.get("use_mammoth_html") and MAMMOTH_AVAILABLE:
            rendered_html = self._extract_with_mammoth(file_path, opts)
        else:
            rendered_html = None

        html_parts.append('<div class="docx-content">')

        # Extract headers if requested
        if opts.get("include_headers"):
            for section in doc.sections:
                header = section.header
                if header and header.paragraphs:
                    header_text = "\n".join(p.text for p in header.paragraphs if p.text.strip())
                    if header_text:
                        para_id = f"header_{uuid.uuid4().hex[:8]}"
                        text_parts.append(header_text)
                        text_parts.append("\n\n")

                        html_parts.append(f'<div class="docx-header" data-para-id="{para_id}">')
                        html_parts.append(f'{html.escape(header_text)}')
                        html_parts.append('</div>')

                        mapper.add_mapping(
                            current_offset,
                            current_offset + len(header_text),
                            DocumentCoordinate(
                                paragraph_id=para_id,
                                local_offset=0,
                                section="header",
                            )
                        )
                        current_offset += len(header_text) + 2
                        metadata["sections"].append({"type": "header", "id": para_id})

        # Extract main body
        current_section = None
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text
            if not para_text.strip():
                continue

            para_id = f"p_{i}_{uuid.uuid4().hex[:8]}"

            # Detect heading level
            heading_level = None
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if style_name.startswith("heading"):
                    try:
                        heading_level = int(style_name.replace("heading", "").strip())
                    except ValueError:
                        pass

            # Update section tracking
            if heading_level:
                current_section = para_text.strip()

            # Build text
            start_offset = current_offset
            text_parts.append(para_text)
            end_offset = current_offset + len(para_text)

            # Add paragraph separator
            text_parts.append(opts["paragraph_separator"])
            current_offset = end_offset + len(opts["paragraph_separator"])

            # Build HTML
            css_class = "docx-paragraph"
            if heading_level:
                css_class = f"docx-heading docx-h{heading_level}"
                html_tag = f"h{min(heading_level, 6)}"
            else:
                html_tag = "p"

            html_parts.append(
                f'<{html_tag} class="{css_class}" '
                f'data-para-id="{para_id}" '
                f'data-start="{start_offset}" '
                f'data-end="{end_offset}">'
                f'{html.escape(para_text)}'
                f'</{html_tag}>'
            )

            # Add coordinate mapping
            mapper.add_mapping(
                start_offset,
                end_offset,
                DocumentCoordinate(
                    paragraph_id=para_id,
                    local_offset=0,
                    section=current_section,
                    heading_level=heading_level,
                )
            )

            # Track paragraph metadata
            metadata["paragraphs"].append({
                "id": para_id,
                "start": start_offset,
                "end": end_offset,
                "heading_level": heading_level,
                "section": current_section,
                "char_count": len(para_text),
            })

        # Extract tables if requested
        if opts.get("include_tables"):
            for t_idx, table in enumerate(doc.tables):
                table_id = f"table_{t_idx}_{uuid.uuid4().hex[:8]}"
                table_text, table_html = self._extract_table(
                    table, table_id, current_offset
                )

                if table_text:
                    text_parts.append("\n")
                    text_parts.append(table_text)
                    text_parts.append("\n")
                    html_parts.append(table_html)

                    mapper.add_mapping(
                        current_offset,
                        current_offset + len(table_text),
                        DocumentCoordinate(
                            paragraph_id=table_id,
                            section=current_section,
                        )
                    )
                    current_offset += len(table_text) + 2
                    metadata["tables"].append({"id": table_id})

        html_parts.append('</div>')

        full_text = "".join(text_parts)

        # Use mammoth HTML if available, otherwise use our generated HTML
        if rendered_html:
            final_html = rendered_html
        else:
            final_html = "\n".join(html_parts)

        coord_dict = mapper.to_dict()
        coord_dict["get_coords_for_range"] = mapper.get_coords_for_range

        return FormatOutput(
            text=full_text,
            rendered_html=final_html,
            coordinate_map=coord_dict,
            metadata=metadata,
            format_name=self.format_name,
            source_path=str(file_path),
        )

    def _extract_with_mammoth(self, file_path: str, opts: Dict[str, Any]) -> str:
        """
        Use mammoth for rich HTML conversion.
        """
        if not MAMMOTH_AVAILABLE:
            return None

        try:
            with open(file_path, "rb") as f:
                result = mammoth.convert_to_html(f)
                html_content = result.value

                # Wrap in container
                return f'<div class="docx-content docx-mammoth">{html_content}</div>'
        except Exception as e:
            logger.warning(f"Mammoth conversion failed: {e}")
            return None

    def _extract_table(
        self,
        table,
        table_id: str,
        base_offset: int
    ) -> tuple:
        """
        Extract text and HTML from a table.

        Returns:
            Tuple of (text, html)
        """
        text_rows = []
        html_parts = []

        html_parts.append(f'<table class="docx-table" data-table-id="{table_id}">')

        for row_idx, row in enumerate(table.rows):
            row_texts = []
            html_parts.append('<tr>')

            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_texts.append(cell_text)
                html_parts.append(f'<td>{html.escape(cell_text)}</td>')

            html_parts.append('</tr>')
            text_rows.append("\t".join(row_texts))

        html_parts.append('</table>')

        return "\n".join(text_rows), "\n".join(html_parts)

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract document metadata (author, title, etc.).

        Args:
            file_path: Path to the .docx file

        Returns:
            Dictionary of metadata properties
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required")

        doc = docx.Document(file_path)
        core_props = doc.core_properties

        return {
            "author": core_props.author,
            "title": core_props.title,
            "subject": core_props.subject,
            "keywords": core_props.keywords,
            "created": str(core_props.created) if core_props.created else None,
            "modified": str(core_props.modified) if core_props.modified else None,
            "last_modified_by": core_props.last_modified_by,
            "revision": core_props.revision,
            "category": core_props.category,
            "comments": core_props.comments,
        }
